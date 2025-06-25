from docx import Document
import os

DOC_PATH = "docs/documentation.docx"

def load_or_create_document():
    if os.path.exists(DOC_PATH):
        return Document(DOC_PATH)
    else:
        return Document()

def update_function_doc(document, func_name, new_doc):
    found = False
    for i, p in enumerate(document.paragraphs):
        if p.text.strip().startswith(f"### {func_name}"):
            found = True
            document.paragraphs[i].text = new_doc
            break

    if not found:
        document.add_paragraph(new_doc)

def save_document(document):
    os.makedirs("docs", exist_ok=True)
    document.save(DOC_PATH)
